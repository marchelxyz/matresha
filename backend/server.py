#!/usr/bin/env python3
"""
AI Assistant Backend Server
Supports multiple AI providers: OpenAI, Gemini, Claude, Groq, Mistral, DeepSeek
"""

import os
import json
import time
import base64
import io
from datetime import datetime, timezone
from flask import Flask, request, jsonify, Response, stream_with_context, g
from flask_cors import CORS
from dotenv import load_dotenv
from contextlib import contextmanager
from werkzeug.utils import secure_filename

# Import database
from database import SessionLocal, Chat, Message, get_db

# Import storage
from storage import storage

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app)

# AI Provider Classes
class AIProvider:
    """Base class for AI providers"""
    
    def __init__(self, api_key=None):
        self.api_key = api_key
    
    def generate(self, message, **kwargs):
        raise NotImplementedError
    
    def stream(self, message, **kwargs):
        raise NotImplementedError


class OpenAIProvider(AIProvider):
    """OpenAI GPT-4 Provider"""
    
    def __init__(self, api_key=None):
        super().__init__(api_key or os.getenv('OPENAI_API_KEY'))
        try:
            import openai
            self.client = openai.OpenAI(api_key=self.api_key) if self.api_key else None
        except ImportError:
            self.client = None
    
    def generate(self, message, temperature=0.7, max_tokens=2000, messages=None, **kwargs):
        if not self.client:
            raise ValueError("OpenAI API key not configured")
        
        # Use provided messages history or create new from single message
        if messages:
            message_list = messages
        else:
            message_list = [{"role": "user", "content": message}]
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=message_list,
                temperature=temperature,
                max_tokens=max_tokens
            )
            
            return response.choices[0].message.content
        except Exception as e:
            error_msg = str(e)
            error_type = type(e).__name__
            
            # Handle RateLimitError and quota errors
            if error_type == 'RateLimitError' or 'insufficient_quota' in error_msg.lower() or 'quota' in error_msg.lower():
                raise ValueError("OpenAI API: Превышена квота использования. Пожалуйста, проверьте ваш план и настройки биллинга. Для получения дополнительной информации: https://platform.openai.com/docs/guides/error-codes/api-errors")
            elif error_type == 'AuthenticationError' or 'api key' in error_msg.lower():
                raise ValueError("OpenAI API: Неверный API ключ или ключ не настроен")
            elif 'rate limit' in error_msg.lower():
                raise ValueError("OpenAI API: Превышен лимит запросов. Пожалуйста, подождите немного и попробуйте снова.")
            else:
                raise ValueError(f"OpenAI API ошибка: {error_msg}")
    
    def stream(self, message, temperature=0.7, max_tokens=2000, messages=None, **kwargs):
        if not self.client:
            raise ValueError("OpenAI API key not configured")
        
        # Use provided messages history or create new from single message
        if messages:
            message_list = messages
        else:
            message_list = [{"role": "user", "content": message}]
        
        try:
            stream = self.client.chat.completions.create(
                model="gpt-4o",
                messages=message_list,
                temperature=temperature,
                max_tokens=max_tokens,
                stream=True
            )
            
            for chunk in stream:
                if chunk.choices[0].delta.content:
                    yield chunk.choices[0].delta.content
        except Exception as e:
            error_msg = str(e)
            error_type = type(e).__name__
            
            # Handle RateLimitError and quota errors
            if error_type == 'RateLimitError' or 'insufficient_quota' in error_msg.lower() or 'quota' in error_msg.lower():
                raise ValueError("OpenAI API: Превышена квота использования. Пожалуйста, проверьте ваш план и настройки биллинга. Для получения дополнительной информации: https://platform.openai.com/docs/guides/error-codes/api-errors")
            elif error_type == 'AuthenticationError' or 'api key' in error_msg.lower():
                raise ValueError("OpenAI API: Неверный API ключ или ключ не настроен")
            elif 'rate limit' in error_msg.lower():
                raise ValueError("OpenAI API: Превышен лимит запросов. Пожалуйста, подождите немного и попробуйте снова.")
            else:
                raise ValueError(f"OpenAI API ошибка: {error_msg}")


class GeminiProvider(AIProvider):
    """Google Gemini Provider with ChatSession support"""
    
    def __init__(self, api_key=None):
        super().__init__(api_key or os.getenv('GEMINI_API_KEY'))
        self.model_name = None
        self.model = None
        self.genai = None
        try:
            import google.generativeai as genai
            self.genai = genai
            if self.api_key:
                genai.configure(api_key=self.api_key)
                # Try models in order - newer models first (gemini-pro is deprecated)
                # gemini-2.5-flash is the newest and fastest
                # gemini-1.5-flash is fast and widely available
                # gemini-1.5-pro is more capable
                model_attempts = [
                    ('gemini-2.5-flash', 'gemini-2.5-flash'),
                    ('gemini-1.5-flash', 'gemini-1.5-flash'),
                    ('gemini-1.5-pro', 'gemini-1.5-pro'),
                    ('gemini-pro', 'gemini-pro'),  # Legacy, may not work with v1beta
                ]
                
                # Try to get list of available models from API
                available_models = set()
                try:
                    for model in genai.list_models():
                        if hasattr(model, 'name'):
                            model_name = model.name.split('/')[-1] if '/' in model.name else model.name
                            available_models.add(model_name)
                except Exception as e:
                    print(f"Gemini: Could not list available models: {str(e)}")
                
                for model_display_name, model_id in model_attempts:
                    try:
                        # Check if model is in available models list (if we got it)
                        if available_models and model_id not in available_models:
                            print(f"Gemini: Model {model_display_name} ({model_id}) not in available models list")
                            continue
                        
                        # Test if model is available by trying to create it
                        test_model = genai.GenerativeModel(model_id)
                        self.model = test_model
                        self.model_name = model_id
                        print(f"Gemini: Successfully initialized model {model_display_name} ({model_id})")
                        break
                    except Exception as e:
                        print(f"Gemini: Model {model_display_name} ({model_id}) not available: {str(e)}")
                        continue
                
                if not self.model:
                    print("Gemini: Warning - No compatible model found. Will attempt to initialize on first request.")
                    # Don't set a default model name, let _ensure_model_initialized handle it
            else:
                print("Gemini: API key not configured")
        except ImportError:
            print("Gemini: google-generativeai library not installed")
            self.model = None
            self.genai = None
    
    def _ensure_model_initialized(self, exclude_models=None):
        """Ensure model is initialized, try to initialize if not
        
        Args:
            exclude_models: List of model names to exclude from attempts
        """
        if self.model:
            return True
        
        if not self.genai or not self.api_key:
            return False
        
        exclude_models = exclude_models or []
        
        # Try to initialize model with fallback - newer models first
        model_attempts = [
            ('gemini-2.5-flash', 'gemini-2.5-flash'),
            ('gemini-1.5-flash', 'gemini-1.5-flash'),
            ('gemini-1.5-pro', 'gemini-1.5-pro'),
            ('gemini-pro', 'gemini-pro'),  # Legacy, may not work with v1beta
        ]
        
        # Try to get list of available models from API
        available_models = set()
        try:
            for model in self.genai.list_models():
                if hasattr(model, 'name'):
                    model_name = model.name.split('/')[-1] if '/' in model.name else model.name
                    available_models.add(model_name)
        except Exception as e:
            print(f"Gemini: Could not list available models: {str(e)}")
        
        for model_display_name, model_id in model_attempts:
            if model_id in exclude_models:
                print(f"Gemini: Skipping model {model_display_name} ({model_id}) - already tried")
                continue
            
            # Check if model is in available models list (if we got it)
            if available_models and model_id not in available_models:
                print(f"Gemini: Model {model_display_name} ({model_id}) not in available models list")
                continue
            
            try:
                self.model = self.genai.GenerativeModel(model_id)
                self.model_name = model_id
                print(f"Gemini: Successfully initialized model {model_display_name} ({model_id})")
                return True
            except Exception as e:
                print(f"Gemini: Model {model_display_name} ({model_id}) failed: {str(e)}")
                continue
        
        return False
    
    def _convert_messages_to_gemini_format(self, messages):
        """Convert messages from OpenAI format to Gemini format"""
        if not messages:
            return []
        
        gemini_messages = []
        for msg in messages:
            role = msg.get('role', 'user')
            content = msg.get('content', '')
            
            # Gemini uses 'user' and 'model' roles (not 'assistant')
            if role == 'assistant':
                gemini_role = 'model'
            elif role == 'user':
                gemini_role = 'user'
            else:
                # Skip system messages or convert to user
                gemini_role = 'user'
            
            gemini_messages.append({
                'role': gemini_role,
                'parts': [content]
            })
        
        return gemini_messages
    
    def generate(self, message, temperature=0.7, max_tokens=2000, messages=None, retry_count=0, **kwargs):
        # Ensure model is initialized
        if not self._ensure_model_initialized():
            raise ValueError("Gemini API key not configured or no compatible model available")
        
        # Prevent infinite recursion
        if retry_count > 2:
            raise ValueError("Gemini API: Failed to initialize compatible model after multiple attempts")
        
        try:
            # If we have message history, use ChatSession for proper context
            if messages and len(messages) > 1:
                # Convert messages to Gemini format
                gemini_messages = self._convert_messages_to_gemini_format(messages)
                
                # History is all messages except the last one (which is the current user message)
                history = gemini_messages[:-1] if len(gemini_messages) > 1 else []
                
                # Current user message is the last one
                current_message = gemini_messages[-1] if gemini_messages else None
                user_content = current_message.get('parts', [message])[0] if current_message else message
                
                # Ensure user_content is a string
                if not isinstance(user_content, str):
                    user_content = str(user_content)
                
                # Create a chat session with history
                chat = self.model.start_chat(history=history)
                
                # Generate response with chat context
                response = chat.send_message(
                    user_content,
                    generation_config={
                        "temperature": temperature,
                        "max_output_tokens": max_tokens,
                    }
                )
            else:
                # Single message, no history
                response = self.model.generate_content(
                    message,
                    generation_config={
                        "temperature": temperature,
                        "max_output_tokens": max_tokens,
                    }
                )
            
            # Handle response
            if hasattr(response, 'text'):
                return response.text
            elif hasattr(response, 'candidates') and response.candidates:
                candidate = response.candidates[0]
                if hasattr(candidate, 'content') and hasattr(candidate.content, 'parts'):
                    return ''.join([part.text for part in candidate.content.parts if hasattr(part, 'text')])
                return str(candidate)
            else:
                return str(response)
                
        except Exception as e:
            error_msg = str(e)
            # If model not found (404) or not supported, try to reinitialize with fallback
            if ("404" in error_msg and "not found" in error_msg.lower()) or \
               ("not supported" in error_msg.lower() and "generateContent" in error_msg):
                print(f"Gemini: Model {self.model_name} not available or not supported, trying fallback...")
                # Reset model and try to initialize with fallback
                old_model_name = self.model_name
                self.model = None
                self.model_name = None
                # Exclude the failed model from retry attempts
                if self._ensure_model_initialized(exclude_models=[old_model_name]):
                    if self.model_name and self.model_name != old_model_name:
                        # Retry with new model
                        try:
                            return self.generate(message, temperature, max_tokens, messages, retry_count=retry_count+1, **kwargs)
                        except Exception as retry_e:
                            error_msg = str(retry_e)
                    else:
                        error_msg = f"Gemini API: Failed to initialize fallback model. Tried: {old_model_name}"
                else:
                    error_msg = f"Gemini API: No compatible model available. Tried: {old_model_name}"
            
            # Provide more helpful error messages
            if "API_KEY" in error_msg or "api key" in error_msg.lower():
                raise ValueError("Gemini API key not configured or invalid")
            elif "quota" in error_msg.lower() or "limit" in error_msg.lower():
                raise ValueError("Gemini API quota exceeded. Please check your usage limits.")
            else:
                raise ValueError(f"Gemini API error: {error_msg}")
    
    def stream(self, message, temperature=0.7, max_tokens=2000, messages=None, retry_count=0, **kwargs):
        # Ensure model is initialized
        if not self._ensure_model_initialized():
            raise ValueError("Gemini API key not configured or no compatible model available")
        
        # Prevent infinite recursion
        if retry_count > 2:
            raise ValueError("Gemini API: Failed to initialize compatible model after multiple attempts")
        
        try:
            # If we have message history, use ChatSession for proper context
            if messages and len(messages) > 1:
                # Convert messages to Gemini format
                gemini_messages = self._convert_messages_to_gemini_format(messages)
                
                # History is all messages except the last one (which is the current user message)
                history = gemini_messages[:-1] if len(gemini_messages) > 1 else []
                
                # Current user message is the last one
                current_message = gemini_messages[-1] if gemini_messages else None
                user_content = current_message.get('parts', [message])[0] if current_message else message
                
                # Ensure user_content is a string
                if not isinstance(user_content, str):
                    user_content = str(user_content)
                
                # Create a chat session with history
                chat = self.model.start_chat(history=history)
                
                # Stream response with chat context
                response = chat.send_message(
                    user_content,
                    generation_config={
                        "temperature": temperature,
                        "max_output_tokens": max_tokens,
                    },
                    stream=True
                )
            else:
                # Single message, no history
                response = self.model.generate_content(
                    message,
                    generation_config={
                        "temperature": temperature,
                        "max_output_tokens": max_tokens,
                    },
                    stream=True
                )
            
            # Stream chunks
            for chunk in response:
                if hasattr(chunk, 'text') and chunk.text:
                    yield chunk.text
                elif hasattr(chunk, 'candidates') and chunk.candidates:
                    candidate = chunk.candidates[0]
                    if hasattr(candidate, 'content') and hasattr(candidate.content, 'parts'):
                        for part in candidate.content.parts:
                            if hasattr(part, 'text') and part.text:
                                yield part.text
                
        except Exception as e:
            error_msg = str(e)
            # If model not found (404) or not supported, try to reinitialize with fallback
            if ("404" in error_msg and "not found" in error_msg.lower()) or \
               ("not supported" in error_msg.lower() and "generateContent" in error_msg):
                print(f"Gemini: Model {self.model_name} not available or not supported, trying fallback...")
                # Reset model and try to initialize with fallback
                old_model_name = self.model_name
                self.model = None
                self.model_name = None
                # Exclude the failed model from retry attempts
                if self._ensure_model_initialized(exclude_models=[old_model_name]):
                    if self.model_name and self.model_name != old_model_name:
                        # Retry with new model
                        try:
                            for chunk in self.stream(message, temperature, max_tokens, messages, retry_count=retry_count+1, **kwargs):
                                yield chunk
                            return
                        except Exception as retry_e:
                            error_msg = str(retry_e)
                    else:
                        error_msg = f"Gemini API: Failed to initialize fallback model. Tried: {old_model_name}"
                else:
                    error_msg = f"Gemini API: No compatible model available. Tried: {old_model_name}"
            
            # Provide more helpful error messages
            if "API_KEY" in error_msg or "api key" in error_msg.lower():
                raise ValueError("Gemini API key not configured or invalid")
            elif "quota" in error_msg.lower() or "limit" in error_msg.lower():
                raise ValueError("Gemini API quota exceeded. Please check your usage limits.")
            else:
                raise ValueError(f"Gemini API error: {error_msg}")


class ClaudeProvider(AIProvider):
    """Anthropic Claude Provider"""
    
    def __init__(self, api_key=None):
        super().__init__(api_key or os.getenv('ANTHROPIC_API_KEY'))
        self.client = None
        if not self.api_key:
            return
        
        try:
            from anthropic import Anthropic
            # Initialize Anthropic client - ensure we only pass api_key
            # Some versions may have issues with additional parameters
            self.client = Anthropic(api_key=self.api_key)
        except ImportError:
            print("ERROR: anthropic library not installed. Install with: pip install anthropic")
            self.client = None
        except TypeError as e:
            # Handle TypeError which may occur if unexpected arguments are passed
            error_msg = str(e)
            if 'proxies' in error_msg or 'unexpected keyword argument' in error_msg:
                print(f"ERROR: Anthropic client initialization error - {error_msg}")
                print("INFO: This may be due to a version incompatibility. Trying alternative method...")
                # Try using environment variable instead
                import os as anthropic_os
                original_key = anthropic_os.environ.get('ANTHROPIC_API_KEY')
                try:
                    anthropic_os.environ['ANTHROPIC_API_KEY'] = self.api_key
                    self.client = Anthropic()
                finally:
                    if original_key:
                        anthropic_os.environ['ANTHROPIC_API_KEY'] = original_key
                    elif 'ANTHROPIC_API_KEY' in anthropic_os.environ:
                        del anthropic_os.environ['ANTHROPIC_API_KEY']
            else:
                raise
        except Exception as e:
            print(f"ERROR: Failed to initialize Anthropic client: {str(e)}")
            self.client = None
    
    def generate(self, message, temperature=0.7, max_tokens=2000, messages=None, **kwargs):
        if not self.client:
            raise ValueError("Anthropic API key not configured")
        
        # Use provided messages history or create new from single message
        if messages:
            message_list = messages
        else:
            message_list = [{"role": "user", "content": message}]
        
        response = self.client.messages.create(
            model="claude-3-opus-20240229",
            max_tokens=max_tokens,
            temperature=temperature,
            messages=message_list
        )
        
        return response.content[0].text
    
    def stream(self, message, temperature=0.7, max_tokens=2000, messages=None, **kwargs):
        if not self.client:
            raise ValueError("Anthropic API key not configured")
        
        # Use provided messages history or create new from single message
        if messages:
            message_list = messages
        else:
            message_list = [{"role": "user", "content": message}]
        
        with self.client.messages.stream(
            model="claude-3-opus-20240229",
            max_tokens=max_tokens,
            temperature=temperature,
            messages=message_list
        ) as stream:
            for text in stream.text_stream:
                yield text


def estimate_tokens(text):
    """
    Оценка количества токенов в тексте.
    Использует консервативную оценку: для английского текста ~4 символа на токен,
    для русского и других языков ~2-3 символа на токен.
    Берем более консервативное значение ~2.5 символа на токен для безопасности.
    """
    if not text:
        return 0
    # Более консервативная оценка: учитываем пробелы и знаки препинания
    # Для смешанного контента используем ~2.5 символа на токен (было 3)
    # Это даст более точную оценку и поможет избежать превышения лимитов
    return int(len(text) / 2.5)


def estimate_messages_tokens(messages):
    """
    Оценка количества токенов в списке сообщений.
    Учитывает структуру сообщений (role, content) и добавляет overhead.
    """
    total_tokens = 0
    for msg in messages:
        # Overhead для структуры сообщения (role, форматирование)
        # Увеличиваем overhead для более точной оценки
        total_tokens += 8  # Было 5, увеличиваем для безопасности
        content = msg.get('content', '')
        if isinstance(content, str):
            total_tokens += estimate_tokens(content)
        elif isinstance(content, list):
            # Для мультимодальных сообщений
            for item in content:
                if isinstance(item, dict):
                    text = item.get('text', '')
                    if text:
                        total_tokens += estimate_tokens(text)
    return total_tokens


def extract_requested_tokens_from_error(error_msg):
    """
    Извлекает количество запрашиваемых токенов из сообщения об ошибке Groq API.
    
    Args:
        error_msg: Строка с сообщением об ошибке
        
    Returns:
        Количество запрашиваемых токенов или None, если не удалось извлечь
    """
    import re
    # Ищем паттерн "Requested XXXXX" в сообщении об ошибке
    # Пример: "Requested 18637"
    match = re.search(r'Requested\s+(\d+)', error_msg, re.IGNORECASE)
    if match:
        return int(match.group(1))
    return None


def calculate_safe_limit_from_error(requested_tokens, max_limit=12000):
    """
    Вычисляет безопасный лимит токенов на основе запрошенного количества.
    
    Args:
        requested_tokens: Количество запрашиваемых токенов
        max_limit: Максимальный лимит TPM (по умолчанию 12000 для Groq on_demand)
        
    Returns:
        Безопасный лимит токенов (обычно 70-80% от max_limit)
    """
    # Используем 70% от максимального лимита для безопасности
    # Это оставляет место для ответа и overhead
    safe_limit = int(max_limit * 0.7)
    
    # Если запрошенное количество больше максимального лимита,
    # используем еще более консервативный подход
    if requested_tokens and requested_tokens > max_limit:
        # Разбиваем на части так, чтобы каждая часть была не больше 50% от лимита
        safe_limit = int(max_limit * 0.5)
    
    return safe_limit


def split_large_request(messages, max_tokens=10000):
    """
    Разбивает большой запрос на части, если он превышает лимит токенов.
    
    Args:
        messages: Список сообщений в формате [{"role": "...", "content": "..."}]
        max_tokens: Максимальное количество токенов на запрос (по умолчанию 10000 для безопасности)
    
    Returns:
        Список частей запроса, каждая часть - список сообщений
    """
    if not messages:
        return [messages]
    
    # Оцениваем общее количество токенов
    total_tokens = estimate_messages_tokens(messages)
    
    # Если запрос не превышает лимит, возвращаем как есть
    if total_tokens <= max_tokens:
        return [messages]
    
    # Если превышает, разбиваем на части
    # Берем последнее сообщение пользователя (оно обычно самое большое)
    last_message = messages[-1]
    last_content = last_message.get('content', '')
    
    # Если проблема в истории сообщений, ограничиваем историю
    if len(messages) > 1:
        history_tokens = estimate_messages_tokens(messages[:-1])
        last_message_tokens = estimate_messages_tokens([last_message])
        
        # Если история слишком большая, сокращаем её
        if history_tokens > max_tokens * 0.7:  # История занимает больше 70% лимита
            # Оставляем только последние N сообщений из истории
            remaining_tokens = max_tokens - last_message_tokens - 100  # Оставляем запас
            kept_messages = []
            kept_tokens = 0
            
            # Идем с конца истории и добавляем сообщения пока не превысим лимит
            for msg in reversed(messages[:-1]):
                msg_tokens = estimate_messages_tokens([msg])
                if kept_tokens + msg_tokens <= remaining_tokens:
                    kept_messages.insert(0, msg)
                    kept_tokens += msg_tokens
                else:
                    break
            
            # Если удалось сократить историю, пробуем снова
            if len(kept_messages) < len(messages) - 1:
                new_messages = kept_messages + [last_message]
                if estimate_messages_tokens(new_messages) <= max_tokens:
                    return [new_messages]
    
    # Если последнее сообщение не слишком большое, проверяем общий размер
    if not isinstance(last_content, str) or len(last_content) <= max_tokens * 2.5:
        # Если общий размер все еще превышает лимит, сокращаем историю
        if total_tokens > max_tokens:
            # Пробуем сократить историю еще больше
            if len(messages) > 2:
                # Оставляем только последние 2-3 сообщения
                reduced_messages = messages[-3:] if len(messages) >= 3 else messages
                if estimate_messages_tokens(reduced_messages) <= max_tokens:
                    return [reduced_messages]
                # Если все еще слишком много, оставляем только последнее сообщение
                return [[last_message]]
        return [messages]
    
    # Разбиваем последнее сообщение на части
    parts = []
    # Уменьшаем размер чанка для более консервативного разбиения
    chunk_size = int(max_tokens * 2.5)  # Было max_tokens * 3, уменьшаем до 2.5
    
    # Разбиваем по предложениям/абзацам для более естественного разбиения
    paragraphs = last_content.split('\n\n')
    current_chunk = []
    current_size = 0
    
    for para in paragraphs:
        para_size = len(para)
        if current_size + para_size > chunk_size and current_chunk:
            # Сохраняем текущий чанк
            chunk_content = '\n\n'.join(current_chunk)
            # Используем сокращенную историю если она была определена выше
            base_messages = messages[:-1] if len(messages) > 1 else []
            chunk_messages = base_messages + [{"role": last_message['role'], "content": chunk_content}]
            parts.append(chunk_messages)
            current_chunk = [para]
            current_size = para_size
        else:
            current_chunk.append(para)
            current_size += para_size + 2  # +2 для '\n\n'
    
    # Добавляем последний чанк
    if current_chunk:
        chunk_content = '\n\n'.join(current_chunk)
        base_messages = messages[:-1] if len(messages) > 1 else []
        chunk_messages = base_messages + [{"role": last_message['role'], "content": chunk_content}]
        parts.append(chunk_messages)
    
    # Если не удалось разбить на части (все в одном абзаце), разбиваем по предложениям
    if len(parts) == 0 or (len(parts) == 1 and estimate_messages_tokens(parts[0]) > max_tokens):
        sentences = last_content.split('. ')
        current_chunk = []
        current_size = 0
        
        for sentence in sentences:
            sentence_size = len(sentence)
            if current_size + sentence_size > chunk_size and current_chunk:
                chunk_content = '. '.join(current_chunk) + '.'
                base_messages = messages[:-1] if len(messages) > 1 else []
                chunk_messages = base_messages + [{"role": last_message['role'], "content": chunk_content}]
                # Проверяем размер перед добавлением
                if estimate_messages_tokens(chunk_messages) <= max_tokens:
                    parts.append(chunk_messages)
                current_chunk = [sentence]
                current_size = sentence_size
            else:
                current_chunk.append(sentence)
                current_size += sentence_size + 2
        
        if current_chunk:
            chunk_content = '. '.join(current_chunk)
            if not chunk_content.endswith('.'):
                chunk_content += '.'
            base_messages = messages[:-1] if len(messages) > 1 else []
            chunk_messages = base_messages + [{"role": last_message['role'], "content": chunk_content}]
            # Проверяем размер перед добавлением
            if estimate_messages_tokens(chunk_messages) <= max_tokens:
                parts.append(chunk_messages)
    
    # Если все еще не удалось разбить, используем простое разбиение по символам
    if len(parts) == 0 or (len(parts) == 1 and estimate_messages_tokens(parts[0]) > max_tokens):
        parts = []
        base_messages = messages[:-1] if len(messages) > 1 else []
        # Еще больше уменьшаем размер чанка для финального разбиения
        final_chunk_size = int(max_tokens * 2.0)  # Еще более консервативно
        for i in range(0, len(last_content), final_chunk_size):
            chunk_content = last_content[i:i+final_chunk_size]
            chunk_messages = base_messages + [{"role": last_message['role'], "content": chunk_content}]
            # Проверяем размер перед добавлением
            if estimate_messages_tokens(chunk_messages) <= max_tokens:
                parts.append(chunk_messages)
            else:
                # Если даже один чанк слишком большой, разбиваем еще мельче
                # Разбиваем на части по половине размера
                sub_chunk_size = final_chunk_size // 2
                for j in range(i, min(i + final_chunk_size, len(last_content)), sub_chunk_size):
                    sub_chunk_content = last_content[j:j+sub_chunk_size]
                    sub_chunk_messages = base_messages + [{"role": last_message['role'], "content": sub_chunk_content}]
                    if estimate_messages_tokens(sub_chunk_messages) <= max_tokens:
                        parts.append(sub_chunk_messages)
    
    return parts if parts else [messages]


class GroqProvider(AIProvider):
    """Groq (Llama) Provider"""
    
    # Лимит токенов для Groq on_demand tier
    MAX_TOKENS_LIMIT = 12000
    # Безопасный лимит с запасом (оставляем место для ответа)
    # Очень консервативный подход - используем 50% от максимального лимита для безопасности
    SAFE_TOKENS_LIMIT = 6000  # 50% от 12000 для максимальной безопасности
    
    def __init__(self, api_key=None):
        super().__init__(api_key or os.getenv('GROQ_API_KEY'))
        try:
            from groq import Groq
            # Groq client initialization - updated to work with groq >= 0.36.0
            if self.api_key:
                self.client = Groq(api_key=self.api_key)
            else:
                self.client = None
        except ImportError:
            self.client = None
        except Exception as e:
            # Log the error for debugging
            error_msg = str(e)
            print(f"ERROR: Failed to initialize Groq client: {error_msg}")
            import traceback
            traceback.print_exc()
            self.client = None
    
    def generate(self, message, temperature=0.7, max_tokens=2000, messages=None, **kwargs):
        if not self.client:
            raise ValueError("Groq API key not configured")
        
        # Use provided messages history or create new from single message
        if messages:
            message_list = messages
        else:
            message_list = [{"role": "user", "content": message}]
        
        # Проверяем размер запроса и разбиваем при необходимости
        # Используем более консервативный подход - 60% от максимального лимита для начального разбиения
        initial_limit = int(self.MAX_TOKENS_LIMIT * 0.6)  # 60% от 12000 = 7200
        request_parts = split_large_request(message_list, initial_limit)
        
        if len(request_parts) > 1:
            # Если запрос разбит на части, обрабатываем каждую часть отдельно
            print(f"INFO: Groq request split into {len(request_parts)} parts")
            responses = []
            for i, part in enumerate(request_parts):
                try:
                    response = self.client.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=part,
                        temperature=temperature,
                        max_tokens=max_tokens
                    )
                    responses.append(response.choices[0].message.content)
                except Exception as e:
                    error_msg = str(e)
                    error_type = type(e).__name__
                    # Проверяем код статуса для APIStatusError
                    status_code = None
                    if hasattr(e, 'status_code'):
                        status_code = e.status_code
                    elif hasattr(e, 'response') and hasattr(e.response, 'status_code'):
                        status_code = e.response.status_code
                    
                    # Если ошибка 413 (Request too large), пытаемся разбить еще мельче
                    if status_code == 413 or '413' in error_msg or 'Request too large' in error_msg or 'tokens per minute' in error_msg.lower() or 'TPM' in error_msg:
                        print(f"WARNING: Part {i+1} still too large ({error_msg}), trying smaller chunks")
                        
                        # Извлекаем точное количество запрашиваемых токенов из ошибки
                        requested_tokens = extract_requested_tokens_from_error(error_msg)
                        if requested_tokens:
                            print(f"INFO: Extracted requested tokens from error: {requested_tokens}")
                            # Вычисляем безопасный лимит на основе запрошенного количества
                            smaller_limit = calculate_safe_limit_from_error(requested_tokens, self.MAX_TOKENS_LIMIT)
                            print(f"INFO: Using calculated safe limit: {smaller_limit}")
                        else:
                            # Если не удалось извлечь, используем консервативный подход
                            smaller_limit = int(self.MAX_TOKENS_LIMIT * 0.4)  # 40% от максимального лимита
                            print(f"INFO: Using conservative safe limit: {smaller_limit}")
                        
                        smaller_parts = split_large_request(part, smaller_limit)
                        
                        # Если все еще одна часть, разбиваем еще мельче
                        if len(smaller_parts) == 1:
                            # Используем еще более агрессивное разбиение
                            if requested_tokens:
                                smaller_limit = int(self.MAX_TOKENS_LIMIT * 0.3)  # 30% от максимального лимита
                            else:
                                smaller_limit = int(self.MAX_TOKENS_LIMIT * 0.25)  # 25% от максимального лимита
                            print(f"INFO: Request still too large, using even smaller limit: {smaller_limit}")
                            smaller_parts = split_large_request(part, smaller_limit)
                        
                        for smaller_part in smaller_parts:
                            try:
                                response = self.client.chat.completions.create(
                                    model="llama-3.3-70b-versatile",
                                    messages=smaller_part,
                                    temperature=temperature,
                                    max_tokens=max_tokens
                                )
                                responses.append(response.choices[0].message.content)
                            except Exception as inner_e:
                                inner_error_msg = str(inner_e)
                                inner_status_code = None
                                if hasattr(inner_e, 'status_code'):
                                    inner_status_code = inner_e.status_code
                                elif hasattr(inner_e, 'response') and hasattr(inner_e.response, 'status_code'):
                                    inner_status_code = inner_e.response.status_code
                                
                                # Если все еще ошибка 413, разбиваем еще мельче
                                if inner_status_code == 413 or '413' in inner_error_msg or 'Request too large' in inner_error_msg or 'tokens per minute' in inner_error_msg.lower() or 'TPM' in inner_error_msg:
                                    print(f"WARNING: Sub-part still too large, trying even smaller chunks")
                                    
                                    # Извлекаем точное количество токенов из ошибки подчасти
                                    inner_requested_tokens = extract_requested_tokens_from_error(inner_error_msg)
                                    if inner_requested_tokens:
                                        print(f"INFO: Extracted requested tokens from sub-part error: {inner_requested_tokens}")
                                        tiny_limit = calculate_safe_limit_from_error(inner_requested_tokens, self.MAX_TOKENS_LIMIT)
                                        # Если все еще слишком много, используем еще более консервативный подход
                                        if inner_requested_tokens > self.MAX_TOKENS_LIMIT * 0.5:
                                            tiny_limit = int(self.MAX_TOKENS_LIMIT * 0.2)  # 20% от максимального лимита
                                    else:
                                        # Используем очень консервативный подход
                                        tiny_limit = int(self.MAX_TOKENS_LIMIT * 0.2)  # 20% от максимального лимита
                                    
                                    print(f"INFO: Using tiny limit: {tiny_limit}")
                                    tiny_parts = split_large_request(smaller_part, tiny_limit)
                                    for tiny_part in tiny_parts:
                                        try:
                                            response = self.client.chat.completions.create(
                                                model="llama-3.3-70b-versatile",
                                                messages=tiny_part,
                                                temperature=temperature,
                                                max_tokens=max_tokens
                                            )
                                            responses.append(response.choices[0].message.content)
                                        except Exception as tiny_e:
                                            raise ValueError(f"Groq API ошибка при обработке микрочасти запроса: {str(tiny_e)}")
                                else:
                                    raise ValueError(f"Groq API ошибка при обработке части запроса: {inner_error_msg}")
                    else:
                        raise ValueError(f"Groq API ошибка: {error_msg}")
            
            # Объединяем ответы
            return '\n\n'.join(responses)
        else:
            # Обычная обработка одного запроса
            try:
                response = self.client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=message_list,
                    temperature=temperature,
                    max_tokens=max_tokens
                )
                return response.choices[0].message.content
            except Exception as e:
                error_msg = str(e)
                error_type = type(e).__name__
                # Проверяем код статуса для APIStatusError
                status_code = None
                if hasattr(e, 'status_code'):
                    status_code = e.status_code
                elif hasattr(e, 'response') and hasattr(e.response, 'status_code'):
                    status_code = e.response.status_code
                
                # Если ошибка 413, пытаемся разбить запрос
                if status_code == 413 or '413' in error_msg or 'Request too large' in error_msg or 'tokens per minute' in error_msg.lower() or 'TPM' in error_msg:
                    print(f"INFO: Request too large ({error_msg}), splitting into smaller parts")
                    
                    # Извлекаем точное количество запрашиваемых токенов из ошибки
                    requested_tokens = extract_requested_tokens_from_error(error_msg)
                    if requested_tokens:
                        print(f"INFO: Extracted requested tokens from error: {requested_tokens}")
                        # Вычисляем безопасный лимит на основе запрошенного количества
                        safe_limit = calculate_safe_limit_from_error(requested_tokens, self.MAX_TOKENS_LIMIT)
                        print(f"INFO: Using calculated safe limit: {safe_limit}")
                    else:
                        # Если не удалось извлечь, используем консервативный подход
                        safe_limit = int(self.MAX_TOKENS_LIMIT * 0.5)  # 50% от максимального лимита
                        print(f"INFO: Using conservative safe limit: {safe_limit}")
                    
                    request_parts = split_large_request(message_list, safe_limit)
                    
                    # Если все еще одна часть, разбиваем еще мельче
                    if len(request_parts) == 1:
                        print(f"INFO: Request still too large after first split, trying smaller limit")
                        if requested_tokens:
                            safe_limit = int(self.MAX_TOKENS_LIMIT * 0.3)  # 30% от максимального лимита
                        else:
                            safe_limit = int(self.MAX_TOKENS_LIMIT * 0.25)  # 25% от максимального лимита
                        request_parts = split_large_request(message_list, safe_limit)
                    
                    if len(request_parts) > 1:
                        # Рекурсивно обрабатываем разбитые части
                        return self.generate(message, temperature, max_tokens, messages, **kwargs)
                raise ValueError(f"Groq API ошибка: {error_msg}")
    
    def stream(self, message, temperature=0.7, max_tokens=2000, messages=None, **kwargs):
        if not self.client:
            raise ValueError("Groq API key not configured")
        
        # Use provided messages history or create new from single message
        if messages:
            message_list = messages
        else:
            message_list = [{"role": "user", "content": message}]
        
        # Проверяем размер запроса и разбиваем при необходимости
        # Используем более консервативный подход - 60% от максимального лимита для начального разбиения
        initial_limit = int(self.MAX_TOKENS_LIMIT * 0.6)  # 60% от 12000 = 7200
        request_parts = split_large_request(message_list, initial_limit)
        
        if len(request_parts) > 1:
            # Если запрос разбит на части, обрабатываем каждую часть отдельно
            print(f"INFO: Groq request split into {len(request_parts)} parts for streaming")
            for i, part in enumerate(request_parts):
                try:
                    # Добавляем разделитель между частями (кроме первой)
                    if i > 0:
                        yield f"\n\n--- Часть {i+1} ---\n\n"
                    
                    stream = self.client.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=part,
                        temperature=temperature,
                        max_tokens=max_tokens,
                        stream=True
                    )
                    
                    for chunk in stream:
                        if chunk.choices and len(chunk.choices) > 0:
                            delta = chunk.choices[0].delta
                            if delta and hasattr(delta, 'content') and delta.content:
                                yield delta.content
                                
                except Exception as e:
                    error_msg = str(e)
                    error_type = type(e).__name__
                    # Проверяем код статуса для APIStatusError
                    status_code = None
                    if hasattr(e, 'status_code'):
                        status_code = e.status_code
                    elif hasattr(e, 'response') and hasattr(e.response, 'status_code'):
                        status_code = e.response.status_code
                    
                    # Если ошибка 413 (Request too large), пытаемся разбить еще мельче
                    if status_code == 413 or '413' in error_msg or 'Request too large' in error_msg or 'tokens per minute' in error_msg.lower() or 'TPM' in error_msg:
                        print(f"WARNING: Part {i+1} still too large ({error_msg}), trying smaller chunks")
                        
                        # Извлекаем точное количество запрашиваемых токенов из ошибки
                        requested_tokens = extract_requested_tokens_from_error(error_msg)
                        if requested_tokens:
                            print(f"INFO: Extracted requested tokens from error: {requested_tokens}")
                            # Вычисляем безопасный лимит на основе запрошенного количества
                            smaller_limit = calculate_safe_limit_from_error(requested_tokens, self.MAX_TOKENS_LIMIT)
                            print(f"INFO: Using calculated safe limit: {smaller_limit}")
                        else:
                            # Если не удалось извлечь, используем консервативный подход
                            smaller_limit = int(self.MAX_TOKENS_LIMIT * 0.4)  # 40% от максимального лимита
                            print(f"INFO: Using conservative safe limit: {smaller_limit}")
                        
                        smaller_parts = split_large_request(part, smaller_limit)
                        
                        # Если все еще одна часть, разбиваем еще мельче
                        if len(smaller_parts) == 1:
                            # Используем еще более агрессивное разбиение
                            if requested_tokens:
                                smaller_limit = int(self.MAX_TOKENS_LIMIT * 0.3)  # 30% от максимального лимита
                            else:
                                smaller_limit = int(self.MAX_TOKENS_LIMIT * 0.25)  # 25% от максимального лимита
                            print(f"INFO: Request still too large, using even smaller limit: {smaller_limit}")
                            smaller_parts = split_large_request(part, smaller_limit)
                        
                        for j, smaller_part in enumerate(smaller_parts):
                            if j > 0:
                                yield f"\n\n--- Подчасть {i+1}.{j+1} ---\n\n"
                            try:
                                stream = self.client.chat.completions.create(
                                    model="llama-3.3-70b-versatile",
                                    messages=smaller_part,
                                    temperature=temperature,
                                    max_tokens=max_tokens,
                                    stream=True
                                )
                                for chunk in stream:
                                    if chunk.choices and len(chunk.choices) > 0:
                                        delta = chunk.choices[0].delta
                                        if delta and hasattr(delta, 'content') and delta.content:
                                            yield delta.content
                            except Exception as inner_e:
                                inner_error_msg = str(inner_e)
                                inner_status_code = None
                                if hasattr(inner_e, 'status_code'):
                                    inner_status_code = inner_e.status_code
                                elif hasattr(inner_e, 'response') and hasattr(inner_e.response, 'status_code'):
                                    inner_status_code = inner_e.response.status_code
                                
                                # Если все еще ошибка 413, разбиваем еще мельче
                                if inner_status_code == 413 or '413' in inner_error_msg or 'Request too large' in inner_error_msg or 'tokens per minute' in inner_error_msg.lower() or 'TPM' in inner_error_msg:
                                    print(f"WARNING: Sub-part {i+1}.{j+1} still too large, trying even smaller chunks")
                                    
                                    # Извлекаем точное количество токенов из ошибки подчасти
                                    inner_requested_tokens = extract_requested_tokens_from_error(inner_error_msg)
                                    if inner_requested_tokens:
                                        print(f"INFO: Extracted requested tokens from sub-part error: {inner_requested_tokens}")
                                        tiny_limit = calculate_safe_limit_from_error(inner_requested_tokens, self.MAX_TOKENS_LIMIT)
                                        # Если все еще слишком много, используем еще более консервативный подход
                                        if inner_requested_tokens > self.MAX_TOKENS_LIMIT * 0.5:
                                            tiny_limit = int(self.MAX_TOKENS_LIMIT * 0.2)  # 20% от максимального лимита
                                    else:
                                        # Используем очень консервативный подход
                                        tiny_limit = int(self.MAX_TOKENS_LIMIT * 0.2)  # 20% от максимального лимита
                                    
                                    print(f"INFO: Using tiny limit: {tiny_limit}")
                                    tiny_parts = split_large_request(smaller_part, tiny_limit)
                                    for k, tiny_part in enumerate(tiny_parts):
                                        if k > 0:
                                            yield f"\n\n--- Микрочасть {i+1}.{j+1}.{k+1} ---\n\n"
                                        try:
                                            stream = self.client.chat.completions.create(
                                                model="llama-3.3-70b-versatile",
                                                messages=tiny_part,
                                                temperature=temperature,
                                                max_tokens=max_tokens,
                                                stream=True
                                            )
                                            for chunk in stream:
                                                if chunk.choices and len(chunk.choices) > 0:
                                                    delta = chunk.choices[0].delta
                                                    if delta and hasattr(delta, 'content') and delta.content:
                                                        yield delta.content
                                        except Exception as tiny_e:
                                            tiny_error_msg = str(tiny_e)
                                            tiny_status_code = None
                                            if hasattr(tiny_e, 'status_code'):
                                                tiny_status_code = tiny_e.status_code
                                            elif hasattr(tiny_e, 'response') and hasattr(tiny_e.response, 'status_code'):
                                                tiny_status_code = tiny_e.response.status_code
                                            
                                            # Если все еще ошибка 413, пытаемся разбить на еще более мелкие части
                                            if tiny_status_code == 413 or '413' in tiny_error_msg or 'Request too large' in tiny_error_msg or 'tokens per minute' in tiny_error_msg.lower() or 'TPM' in tiny_error_msg:
                                                print(f"WARNING: Micro-part {i+1}.{j+1}.{k+1} still too large, trying ultra-small chunks")
                                                
                                                # Извлекаем точное количество токенов из ошибки микрочасти
                                                tiny_requested_tokens = extract_requested_tokens_from_error(tiny_error_msg)
                                                if tiny_requested_tokens:
                                                    print(f"INFO: Extracted requested tokens from micro-part error: {tiny_requested_tokens}")
                                                    ultra_limit = calculate_safe_limit_from_error(tiny_requested_tokens, self.MAX_TOKENS_LIMIT)
                                                    # Используем очень консервативный подход - 15% от максимального лимита
                                                    if tiny_requested_tokens > self.MAX_TOKENS_LIMIT * 0.3:
                                                        ultra_limit = int(self.MAX_TOKENS_LIMIT * 0.15)  # 15% от максимального лимита
                                                else:
                                                    # Используем очень консервативный подход
                                                    ultra_limit = int(self.MAX_TOKENS_LIMIT * 0.15)  # 15% от максимального лимита
                                                
                                                print(f"INFO: Using ultra limit: {ultra_limit}")
                                                ultra_parts = split_large_request(tiny_part, ultra_limit)
                                                
                                                for l, ultra_part in enumerate(ultra_parts):
                                                    if l > 0:
                                                        yield f"\n\n--- Ультра-часть {i+1}.{j+1}.{k+1}.{l+1} ---\n\n"
                                                    try:
                                                        stream = self.client.chat.completions.create(
                                                            model="llama-3.3-70b-versatile",
                                                            messages=ultra_part,
                                                            temperature=temperature,
                                                            max_tokens=max_tokens,
                                                            stream=True
                                                        )
                                                        for chunk in stream:
                                                            if chunk.choices and len(chunk.choices) > 0:
                                                                delta = chunk.choices[0].delta
                                                                if delta and hasattr(delta, 'content') and delta.content:
                                                                    yield delta.content
                                                    except Exception as ultra_e:
                                                        # Если даже ультра-часть не работает, выдаем ошибку
                                                        raise ValueError(f"Groq API ошибка при обработке ультра-части запроса: {str(ultra_e)}")
                                            else:
                                                raise ValueError(f"Groq API ошибка при обработке микрочасти запроса: {tiny_error_msg}")
                                else:
                                    raise ValueError(f"Groq API ошибка при обработке части запроса: {inner_error_msg}")
                    else:
                        raise ValueError(f"Groq API ошибка: {error_msg}")
        else:
            # Обычная обработка одного запроса
            try:
                stream = self.client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=message_list,
                    temperature=temperature,
                    max_tokens=max_tokens,
                    stream=True
                )
                
                for chunk in stream:
                    if chunk.choices and len(chunk.choices) > 0:
                        delta = chunk.choices[0].delta
                        if delta and hasattr(delta, 'content') and delta.content:
                            yield delta.content
            except Exception as e:
                error_msg = str(e)
                error_type = type(e).__name__
                # Проверяем код статуса для APIStatusError
                status_code = None
                if hasattr(e, 'status_code'):
                    status_code = e.status_code
                elif hasattr(e, 'response') and hasattr(e.response, 'status_code'):
                    status_code = e.response.status_code
                
                # Если ошибка 413, пытаемся разбить запрос
                if status_code == 413 or '413' in error_msg or 'Request too large' in error_msg or 'tokens per minute' in error_msg.lower() or 'TPM' in error_msg:
                    print(f"INFO: Request too large ({error_msg}), splitting into smaller parts")
                    
                    # Извлекаем точное количество запрашиваемых токенов из ошибки
                    requested_tokens = extract_requested_tokens_from_error(error_msg)
                    if requested_tokens:
                        print(f"INFO: Extracted requested tokens from error: {requested_tokens}")
                        # Вычисляем безопасный лимит на основе запрошенного количества
                        safe_limit = calculate_safe_limit_from_error(requested_tokens, self.MAX_TOKENS_LIMIT)
                        print(f"INFO: Using calculated safe limit: {safe_limit}")
                    else:
                        # Если не удалось извлечь, используем консервативный подход
                        safe_limit = int(self.MAX_TOKENS_LIMIT * 0.5)  # 50% от максимального лимита
                        print(f"INFO: Using conservative safe limit: {safe_limit}")
                    
                    request_parts = split_large_request(message_list, safe_limit)
                    
                    # Если все еще одна часть, разбиваем еще мельче
                    if len(request_parts) == 1:
                        print(f"INFO: Request still too large after first split, trying smaller limit")
                        if requested_tokens:
                            safe_limit = int(self.MAX_TOKENS_LIMIT * 0.3)  # 30% от максимального лимита
                        else:
                            safe_limit = int(self.MAX_TOKENS_LIMIT * 0.25)  # 25% от максимального лимита
                        request_parts = split_large_request(message_list, safe_limit)
                    
                    if len(request_parts) > 1:
                        # Обрабатываем каждую часть отдельно
                        for i, part in enumerate(request_parts):
                            if i > 0:
                                yield f"\n\n--- Часть {i+1} ---\n\n"
                            try:
                                stream = self.client.chat.completions.create(
                                    model="llama-3.3-70b-versatile",
                                    messages=part,
                                    temperature=temperature,
                                    max_tokens=max_tokens,
                                    stream=True
                                )
                                for chunk in stream:
                                    if chunk.choices and len(chunk.choices) > 0:
                                        delta = chunk.choices[0].delta
                                        if delta and hasattr(delta, 'content') and delta.content:
                                            yield delta.content
                            except Exception as part_e:
                                part_error_msg = str(part_e)
                                part_status_code = None
                                if hasattr(part_e, 'status_code'):
                                    part_status_code = part_e.status_code
                                elif hasattr(part_e, 'response') and hasattr(part_e.response, 'status_code'):
                                    part_status_code = part_e.response.status_code
                                
                                # Если все еще ошибка 413, разбиваем еще мельче
                                if part_status_code == 413 or '413' in part_error_msg or 'Request too large' in part_error_msg or 'tokens per minute' in part_error_msg.lower() or 'TPM' in part_error_msg:
                                    print(f"WARNING: Part {i+1} still too large, trying even smaller chunks")
                                    
                                    # Извлекаем точное количество токенов из ошибки части
                                    part_requested_tokens = extract_requested_tokens_from_error(part_error_msg)
                                    if part_requested_tokens:
                                        print(f"INFO: Extracted requested tokens from part error: {part_requested_tokens}")
                                        tiny_limit = calculate_safe_limit_from_error(part_requested_tokens, self.MAX_TOKENS_LIMIT)
                                        if part_requested_tokens > self.MAX_TOKENS_LIMIT * 0.5:
                                            tiny_limit = int(self.MAX_TOKENS_LIMIT * 0.2)  # 20% от максимального лимита
                                    else:
                                        tiny_limit = int(self.MAX_TOKENS_LIMIT * 0.2)  # 20% от максимального лимита
                                    
                                    print(f"INFO: Using tiny limit: {tiny_limit}")
                                    tiny_parts = split_large_request(part, tiny_limit)
                                    for j, tiny_part in enumerate(tiny_parts):
                                        if j > 0:
                                            yield f"\n\n--- Подчасть {i+1}.{j+1} ---\n\n"
                                        try:
                                            stream = self.client.chat.completions.create(
                                                model="llama-3.3-70b-versatile",
                                                messages=tiny_part,
                                                temperature=temperature,
                                                max_tokens=max_tokens,
                                                stream=True
                                            )
                                            for chunk in stream:
                                                if chunk.choices and len(chunk.choices) > 0:
                                                    delta = chunk.choices[0].delta
                                                    if delta and hasattr(delta, 'content') and delta.content:
                                                        yield delta.content
                                        except Exception as tiny_e:
                                            raise ValueError(f"Groq API ошибка при обработке подчасти запроса: {str(tiny_e)}")
                                else:
                                    raise ValueError(f"Groq API ошибка при обработке части запроса: {part_error_msg}")
                        return
                raise ValueError(f"Groq API ошибка: {error_msg}")


class MistralProvider(AIProvider):
    """Mistral AI Provider (using OpenAI-compatible API)"""
    
    def __init__(self, api_key=None):
        super().__init__(api_key or os.getenv('MISTRAL_API_KEY'))
        try:
            import openai
            if self.api_key:
                self.client = openai.OpenAI(
                    api_key=self.api_key,
                    base_url="https://api.mistral.ai/v1"
                )
            else:
                self.client = None
        except ImportError:
            self.client = None
    
    def generate(self, message, temperature=0.7, max_tokens=2000, messages=None, **kwargs):
        if not self.client:
            raise ValueError("Mistral API key not configured")
        
        # Use provided messages history or create new from single message
        if messages:
            message_list = messages
        else:
            message_list = [{"role": "user", "content": message}]
        
        response = self.client.chat.completions.create(
            model="mistral-large-latest",
            messages=message_list,
            temperature=temperature,
            max_tokens=max_tokens
        )
        
        return response.choices[0].message.content
    
    def stream(self, message, temperature=0.7, max_tokens=2000, messages=None, **kwargs):
        if not self.client:
            raise ValueError("Mistral API key not configured")
        
        # Use provided messages history or create new from single message
        if messages:
            message_list = messages
        else:
            message_list = [{"role": "user", "content": message}]
        
        stream = self.client.chat.completions.create(
            model="mistral-large-latest",
            messages=message_list,
            temperature=temperature,
            max_tokens=max_tokens,
            stream=True
        )
        
        for chunk in stream:
            if chunk.choices[0].delta.content:
                yield chunk.choices[0].delta.content


class DeepSeekProvider(AIProvider):
    """DeepSeek AI Provider (using OpenAI-compatible API)"""
    
    def __init__(self, api_key=None):
        super().__init__(api_key or os.getenv('DEEPSEEK_API_KEY'))
        try:
            import openai
            if self.api_key:
                self.client = openai.OpenAI(
                    api_key=self.api_key,
                    base_url="https://api.deepseek.com/v1"
                )
            else:
                self.client = None
        except ImportError:
            self.client = None
    
    def generate(self, message, temperature=0.7, max_tokens=2000, messages=None, **kwargs):
        if not self.client:
            raise ValueError("DeepSeek API key not configured")
        
        # Use provided messages history or create new from single message
        if messages:
            message_list = messages
        else:
            message_list = [{"role": "user", "content": message}]
        
        try:
            response = self.client.chat.completions.create(
                model="deepseek-chat",
                messages=message_list,
                temperature=temperature,
                max_tokens=max_tokens
            )
            
            return response.choices[0].message.content
        except Exception as e:
            error_msg = str(e)
            error_type = type(e).__name__
            
            # Проверяем код ошибки 402 и сообщение "Insufficient Balance"
            if '402' in error_msg or 'insufficient balance' in error_msg.lower() or 'insufficient_balance' in error_msg.lower():
                raise ValueError("DeepSeek API: Недостаточно средств на балансе. Пожалуйста, пополните ваш аккаунт на https://platform.deepseek.com")
            # Handle API errors
            elif error_type == 'AuthenticationError' or 'api key' in error_msg.lower():
                raise ValueError("DeepSeek API: Неверный API ключ или ключ не настроен")
            elif 'rate limit' in error_msg.lower():
                raise ValueError("DeepSeek API: Превышен лимит запросов. Пожалуйста, подождите немного и попробуйте снова.")
            elif 'quota' in error_msg.lower():
                raise ValueError("DeepSeek API: Превышена квота использования. Пожалуйста, проверьте ваш план.")
            else:
                raise ValueError(f"DeepSeek API ошибка: {error_msg}")
    
    def stream(self, message, temperature=0.7, max_tokens=2000, messages=None, **kwargs):
        if not self.client:
            raise ValueError("DeepSeek API key not configured")
        
        # Use provided messages history or create new from single message
        if messages:
            message_list = messages
        else:
            message_list = [{"role": "user", "content": message}]
        
        try:
            stream = self.client.chat.completions.create(
                model="deepseek-chat",
                messages=message_list,
                temperature=temperature,
                max_tokens=max_tokens,
                stream=True
            )
            
            for chunk in stream:
                if chunk.choices and len(chunk.choices) > 0:
                    delta = chunk.choices[0].delta
                    if delta and hasattr(delta, 'content') and delta.content:
                        yield delta.content
        except Exception as e:
            error_msg = str(e)
            error_type = type(e).__name__
            
            # Проверяем код ошибки 402 и сообщение "Insufficient Balance"
            if '402' in error_msg or 'insufficient balance' in error_msg.lower() or 'insufficient_balance' in error_msg.lower():
                raise ValueError("DeepSeek API: Недостаточно средств на балансе. Пожалуйста, пополните ваш аккаунт на https://platform.deepseek.com")
            # Handle API errors
            elif error_type == 'AuthenticationError' or 'api key' in error_msg.lower():
                raise ValueError("DeepSeek API: Неверный API ключ или ключ не настроен")
            elif 'rate limit' in error_msg.lower():
                raise ValueError("DeepSeek API: Превышен лимит запросов. Пожалуйста, подождите немного и попробуйте снова.")
            elif 'quota' in error_msg.lower():
                raise ValueError("DeepSeek API: Превышена квота использования. Пожалуйста, проверьте ваш план.")
            else:
                raise ValueError(f"DeepSeek API ошибка: {error_msg}")


# Provider registry
PROVIDERS = {
    'openai': OpenAIProvider,
    'gemini': GeminiProvider,
    'claude': ClaudeProvider,
    'groq': GroqProvider,
    'mistral': MistralProvider,
    'deepseek': DeepSeekProvider
}


def get_provider(provider_name):
    """Get provider instance"""
    provider_class = PROVIDERS.get(provider_name)
    if not provider_class:
        raise ValueError(f"Unknown provider: {provider_name}")
    
    provider = provider_class()
    
    # Check if API key is configured
    if not provider.api_key:
        env_var_name = {
            'openai': 'OPENAI_API_KEY',
            'gemini': 'GEMINI_API_KEY',
            'claude': 'ANTHROPIC_API_KEY',
            'groq': 'GROQ_API_KEY',
            'mistral': 'MISTRAL_API_KEY',
            'deepseek': 'DEEPSEEK_API_KEY'
        }.get(provider_name, f'{provider_name.upper()}_API_KEY')
        
        raise ValueError(
            f"{provider_name.capitalize()} API key not configured. "
            f"Please set {env_var_name} environment variable."
        )
    
    return provider


# Database helper functions
def get_or_create_chat(user_id, user_name=None, username=None, provider='openai'):
    """Get existing chat or create new one for user. Returns chat_id (int)."""
    db = next(get_db())
    try:
        # Try to get the most recent chat for this user
        chat = db.query(Chat).filter(
            Chat.user_id == str(user_id),
            Chat.provider == provider
        ).order_by(Chat.updated_at.desc()).first()
        
        if not chat:
            # Create new chat
            chat = Chat(
                user_id=str(user_id),
                user_name=user_name,
                username=username,
                provider=provider
            )
            db.add(chat)
            db.commit()
            db.refresh(chat)
        else:
            # Update chat timestamp
            chat.updated_at = datetime.now(timezone.utc)
            db.commit()
        
        # Сохраняем chat_id до закрытия сессии, чтобы избежать DetachedInstanceError
        chat_id = chat.id
        return chat_id
    finally:
        db.close()


def save_message(chat_id, role, content, provider=None, temperature=None, max_tokens=None):
    """Save message to database"""
    db = next(get_db())
    try:
        message = Message(
            chat_id=chat_id,
            role=role,
            content=content,
            provider=provider,
            temperature=str(temperature) if temperature else None,
            max_tokens=max_tokens
        )
        db.add(message)
        db.commit()
        db.refresh(message)
        return message
    finally:
        db.close()


def get_chat_history(chat_id, limit=50):
    """Get chat history messages"""
    db = next(get_db())
    try:
        messages = db.query(Message).filter(
            Message.chat_id == chat_id
        ).order_by(Message.created_at.asc()).limit(limit).all()
        
        return [msg.to_dict() for msg in messages]
    finally:
        db.close()


def get_user_chats(user_id, limit=10):
    """Get user's chat sessions"""
    db = next(get_db())
    try:
        chats = db.query(Chat).filter(
            Chat.user_id == str(user_id)
        ).order_by(Chat.updated_at.desc()).limit(limit).all()
        
        return [chat.to_dict() for chat in chats]
    finally:
        db.close()


def get_chat_with_messages(chat_id):
    """Get chat with all messages"""
    db = next(get_db())
    try:
        chat = db.query(Chat).filter(Chat.id == chat_id).first()
        if chat:
            chat_dict = chat.to_dict()
            chat_dict['messages'] = get_chat_history(chat_id)
            return chat_dict
        return None
    finally:
        db.close()


@app.route('/api/health', methods=['GET'])
@app.route('/health', methods=['GET'])
def health():
    """Health check endpoint"""
    return jsonify({
        'status': 'ok',
        'timestamp': time.time(),
        'version': '1.0.0'
    })


@app.route('/api/debug/env', methods=['GET'])
@app.route('/debug/env', methods=['GET'])
def debug_env():
    """Debug endpoint to check environment variables (without exposing keys)"""
    env_vars = {}
    for key in os.environ.keys():
        if 'API_KEY' in key or 'KEY' in key:
            value = os.environ.get(key)
            env_vars[key] = {
                'exists': True,
                'has_value': bool(value),
                'length': len(value) if value else 0,
                'starts_with': value[:4] if value and len(value) >= 4 else None,
                'ends_with': value[-4:] if value and len(value) >= 4 else None
            }
    
    return jsonify({
        'success': True,
        'data': {
            'env_vars': env_vars,
            'total_env_vars': len(os.environ),
            'api_key_vars_found': len([k for k in os.environ.keys() if 'API_KEY' in k])
        }
    })


@app.route('/api/providers', methods=['GET'])
@app.route('/providers', methods=['GET'])
def get_providers():
    """Get list of available providers"""
    available = []
    provider_status = {}
    env_vars_status = {}
    
    # Check environment variables directly
    env_var_map = {
        'openai': 'OPENAI_API_KEY',
        'gemini': 'GEMINI_API_KEY',
        'claude': 'ANTHROPIC_API_KEY',
        'groq': 'GROQ_API_KEY',
        'mistral': 'MISTRAL_API_KEY',
        'deepseek': 'DEEPSEEK_API_KEY'
    }
    
    for name, env_var in env_var_map.items():
        env_value = os.getenv(env_var)
        env_vars_status[env_var] = {
            'exists': env_value is not None,
            'has_value': bool(env_value),
            'length': len(env_value) if env_value else 0,
            'starts_with': env_value[:4] if env_value and len(env_value) >= 4 else None
        }
    
    for name, provider_class in PROVIDERS.items():
        try:
            provider = provider_class()
            has_key = bool(provider.api_key)
            provider_status[name] = {
                'available': has_key,
                'has_key': has_key,
                'env_var': env_var_map.get(name),
                'env_var_exists': os.getenv(env_var_map.get(name)) is not None
            }
            if has_key:
                available.append(name)
        except Exception as e:
            provider_status[name] = {
                'available': False,
                'error': str(e),
                'env_var': env_var_map.get(name),
                'env_var_exists': os.getenv(env_var_map.get(name)) is not None
            }
    
    return jsonify({
        'success': True,
        'data': {
            'providers': available,
            'all': list(PROVIDERS.keys()),
            'status': provider_status,
            'env_vars': env_vars_status,
            'debug': {
                'python_env_keys': [k for k in os.environ.keys() if 'API_KEY' in k or 'KEY' in k]
            }
        }
    })


@app.route('/api/chat', methods=['POST'])
@app.route('/chat', methods=['POST'])
def chat():
    """Non-streaming chat endpoint"""
    try:
        data = request.get_json()
        message = data.get('message', '')
        provider_name = data.get('provider', 'openai')
        temperature = float(data.get('temperature', 0.7))
        max_tokens = int(data.get('maxTokens', 2000))
        
        # Get user info
        user_data = data.get('user', {})
        user_id = user_data.get('id') if user_data else None
        user_name = user_data.get('first_name') if user_data else None
        username = user_data.get('username') if user_data else None
        
        if not message:
            return jsonify({
                'success': False,
                'error': 'Message is required'
            }), 400
        
        # Get or create chat session
        chat_id = None
        messages_history = []
        if user_id:
            chat_id = get_or_create_chat(user_id, user_name, username, provider_name)
            # Load chat history
            history = get_chat_history(chat_id)
            # Convert to format expected by providers
            messages_history = [
                {"role": msg['role'], "content": msg['content']}
                for msg in history
            ]
        
        provider = get_provider(provider_name)
        
        # Add current user message to history
        messages_history.append({"role": "user", "content": message})
        
        response = provider.generate(
            message,
            temperature=temperature,
            max_tokens=max_tokens,
            messages=messages_history if messages_history else None
        )
        
        # Save messages to database
        if chat_id and user_id:
            save_message(chat_id, "user", message, provider_name, temperature, max_tokens)
            save_message(chat_id, "assistant", response, provider_name, temperature, max_tokens)
        
        return jsonify({
            'success': True,
            'data': {
                'response': response,
                'provider': provider_name,
                'chat_id': chat_id
            }
        })
    
    except ValueError as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 400
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/chat/stream', methods=['POST'])
@app.route('/chat/stream', methods=['POST'])
def chat_stream():
    """Streaming chat endpoint"""
    try:
        data = request.get_json()
        
        # Логирование для отладки
        if not data:
            print("ERROR: Request body is None or empty")
            return jsonify({
                'success': False,
                'error': 'Request body is required (JSON)'
            }), 400
        
        # Логируем полученные данные (без чувствительной информации)
        print(f"DEBUG: Received stream request - provider: {data.get('provider')}, "
              f"has_message: {bool(data.get('message'))}, "
              f"message_length: {len(data.get('message', ''))}")
        
        message = data.get('message', '')
        provider_name = data.get('provider', 'openai')
        
        # Get user info
        user_data = data.get('user', {})
        user_id = user_data.get('id') if user_data else None
        user_name = user_data.get('first_name') if user_data else None
        username = user_data.get('username') if user_data else None
        
        # Безопасное получение параметров с валидацией
        try:
            temperature = float(data.get('temperature', 0.7))
        except (ValueError, TypeError):
            temperature = 0.7
        
        try:
            max_tokens = int(data.get('maxTokens', 2000))
        except (ValueError, TypeError):
            max_tokens = 2000
        
        if not message:
            return jsonify({
                'success': False,
                'error': 'Message is required'
            }), 400
        
        if not isinstance(message, str) or len(message.strip()) == 0:
            print(f"ERROR: Invalid message - type: {type(message)}, value: {repr(message)}")
            return jsonify({
                'success': False,
                'error': 'Message must be a non-empty string'
            }), 400
        
        try:
            provider = get_provider(provider_name)
        except ValueError as e:
            print(f"ERROR: Provider error - {str(e)}")
            return jsonify({
                'success': False,
                'error': str(e)
            }), 400
        
        # Get or create chat session
        chat_id = None
        messages_history = []
        if user_id:
            chat_id = get_or_create_chat(user_id, user_name, username, provider_name)
            # Load chat history
            history = get_chat_history(chat_id)
            # Convert to format expected by providers
            messages_history = [
                {"role": msg['role'], "content": msg['content']}
                for msg in history
            ]
        
        # Add current user message to history
        messages_history.append({"role": "user", "content": message})
        
        # Save user message to database
        if chat_id and user_id:
            try:
                save_message(chat_id, "user", message, provider_name, temperature, max_tokens)
            except Exception as db_error:
                print(f"WARNING: Failed to save user message to database: {str(db_error)}")
                # Продолжаем выполнение даже если сохранение не удалось
        
        def generate():
            full_response = ""
            chunk_count = 0
            try:
                print(f"DEBUG: Starting stream for provider {provider_name}, message length: {len(message)}")
                for chunk in provider.stream(
                    message,
                    temperature=temperature,
                    max_tokens=max_tokens,
                    messages=messages_history if messages_history else None
                ):
                    if chunk:  # Проверяем, что chunk не пустой
                        chunk_count += 1
                        full_response += chunk
                        try:
                            yield f"data: {json.dumps({'content': chunk})}\n\n"
                        except Exception as yield_error:
                            print(f"ERROR: Failed to yield chunk {chunk_count}: {str(yield_error)}")
                            # Продолжаем обработку даже если не удалось отправить один чанк
                            continue
                
                print(f"DEBUG: Stream completed. Total chunks: {chunk_count}, total length: {len(full_response)}")
                
                # Save assistant response to database
                if chat_id and user_id and full_response:
                    try:
                        save_message(chat_id, "assistant", full_response, provider_name, temperature, max_tokens)
                    except Exception as db_error:
                        print(f"WARNING: Failed to save message to database: {str(db_error)}")
                
                yield "data: [DONE]\n\n"
            except ValueError as e:
                # Handle ValueError (API errors from providers)
                error_msg = str(e)
                error_type = type(e).__name__
                print(f"ERROR: Stream generation failed - {error_type}: {error_msg}")
                # Send user-friendly error message (already formatted by provider)
                yield f"data: {json.dumps({'error': error_msg, 'type': error_type})}\n\n"
                yield "data: [DONE]\n\n"
            except GeneratorExit:
                # Клиент закрыл соединение, это нормально
                print(f"DEBUG: Client closed connection. Processed {chunk_count} chunks, total length: {len(full_response)}")
                raise
            except Exception as e:
                # Log the full error for debugging
                error_msg = str(e)
                error_type = type(e).__name__
                print(f"ERROR: Stream generation failed - {error_type}: {error_msg}")
                print(f"DEBUG: Processed {chunk_count} chunks before error, total length: {len(full_response)}")
                import traceback
                traceback.print_exc()
                
                # Если уже была отправлена часть ответа, сохраняем её
                if chat_id and user_id and full_response:
                    try:
                        save_message(chat_id, "assistant", full_response, provider_name, temperature, max_tokens)
                    except Exception as db_error:
                        print(f"WARNING: Failed to save partial message to database: {str(db_error)}")
                
                # Provide user-friendly message for quota/balance errors
                if 'insufficient balance' in error_msg.lower() or 'insufficient_balance' in error_msg.lower() or '402' in error_msg:
                    user_msg = "DeepSeek API: Недостаточно средств на балансе. Пожалуйста, пополните ваш аккаунт на https://platform.deepseek.com"
                elif '413' in error_msg or 'Request too large' in error_msg.lower() or 'tokens per minute' in error_msg.lower():
                    # Ошибка 413 от Groq - запрос слишком большой
                    user_msg = "Запрос слишком большой. Система автоматически разбивает его на части. Если ошибка повторяется, попробуйте сократить размер сообщения или истории."
                elif 'quota' in error_msg.lower() or 'insufficient_quota' in error_msg.lower() or error_type == 'RateLimitError':
                    user_msg = "Превышена квота использования API. Пожалуйста, проверьте ваш план и настройки биллинга."
                else:
                    user_msg = f"Ошибка при генерации ответа: {error_msg}"
                
                try:
                    yield f"data: {json.dumps({'error': user_msg, 'type': error_type, 'details': error_msg})}\n\n"
                    yield "data: [DONE]\n\n"
                except Exception as yield_error:
                    print(f"ERROR: Failed to yield error message: {str(yield_error)}")
                    # Если не удалось отправить ошибку, просто завершаем
                    pass
        
        return Response(
            stream_with_context(generate()),
            mimetype='text/event-stream',
            headers={
                'Cache-Control': 'no-cache',
                'X-Accel-Buffering': 'no'
            }
        )
    
    except ValueError as e:
        print(f"ERROR: ValueError in chat_stream - {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 400
    except Exception as e:
        error_type = type(e).__name__
        error_msg = str(e)
        print(f"ERROR: Exception in chat_stream - {error_type}: {error_msg}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': error_msg,
            'type': error_type
        }), 500


def process_file(file):
    """Process uploaded file and return content description"""
    try:
        filename = secure_filename(file.filename)
        content_type = file.content_type or ''
        file_extension = filename.lower().split('.')[-1] if '.' in filename else ''
        
        # Read file data once (file.read() moves the file pointer)
        file.seek(0)  # Reset file pointer to beginning
        file_data = file.read()
        
        # Сохраняем файл в хранилище (если настроено)
        storage_info = None
        if storage.enabled:
            storage_info = storage.upload_file(
                file_data=file_data,
                filename=filename,
                content_type=content_type,
                folder='chat-files'
            )
        
        # Process images
        if content_type.startswith('image/'):
            file_base64 = base64.b64encode(file_data).decode('utf-8')
            
            # Try to extract image metadata using Pillow
            image_info = f'Изображение: {filename}'
            try:
                from PIL import Image
                img = Image.open(io.BytesIO(file_data))
                width, height = img.size
                image_info += f' ({width}x{height}px)'
            except Exception:
                pass
            
            result = {
                'type': 'image',
                'filename': filename,
                'content_type': content_type,
                'base64': file_base64,
                'description': image_info
            }
            
            # Добавляем информацию о хранилище, если файл был сохранен
            if storage_info:
                result['storage'] = storage_info
                result['description'] += f'\nФайл сохранен в облаке: {storage_info["url"]}'
            
            return result
        
        # Process PDF files
        elif content_type == 'application/pdf' or file_extension == 'pdf':
            try:
                import pdfplumber
                
                # Extract text from PDF
                text_parts = []
                with pdfplumber.open(io.BytesIO(file_data)) as pdf:
                    for i, page in enumerate(pdf.pages[:10]):  # Limit to first 10 pages
                        text = page.extract_text()
                        if text:
                            text_parts.append(f'Страница {i+1}:\n{text}')
                
                if text_parts:
                    full_text = '\n\n'.join(text_parts)
                    # Limit text length
                    if len(full_text) > 5000:
                        full_text = full_text[:5000] + '\n\n[... текст обрезан ...]'
                    result = {
                        'type': 'pdf',
                        'filename': filename,
                        'content': full_text,
                        'description': f'PDF файл {filename}:\n{full_text}'
                    }
                    
                    # Добавляем информацию о хранилище, если файл был сохранен
                    if storage_info:
                        result['storage'] = storage_info
                        result['description'] += f'\n\nФайл сохранен в облаке: {storage_info["url"]}'
                    
                    return result
                else:
                    result = {
                        'type': 'pdf',
                        'filename': filename,
                        'description': f'PDF файл: {filename} (не удалось извлечь текст, возможно это сканированное изображение)'
                    }
                    if storage_info:
                        result['storage'] = storage_info
                        result['description'] += f'\n\nФайл сохранен в облаке: {storage_info["url"]}'
                    return result
            except ImportError:
                return {
                    'type': 'pdf',
                    'filename': filename,
                    'description': f'PDF файл: {filename} (библиотека pdfplumber не установлена)'
                }
            except Exception as e:
                return {
                    'type': 'pdf',
                    'filename': filename,
                    'description': f'PDF файл: {filename} (ошибка обработки: {str(e)})'
                }
        
        # Process Word documents (.docx)
        elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or file_extension == 'docx':
            try:
                import docx
                
                # Extract text from Word document
                doc = docx.Document(io.BytesIO(file_data))
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                
                full_text = '\n'.join(paragraphs)
                
                # Also extract text from tables
                for table in doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                        table_text.append(row_text)
                    if table_text:
                        full_text += '\n\nТаблица:\n' + '\n'.join(table_text)
                
                if full_text:
                    if len(full_text) > 5000:
                        full_text = full_text[:5000] + '\n\n[... текст обрезан ...]'
                    result = {
                        'type': 'docx',
                        'filename': filename,
                        'content': full_text,
                        'description': f'Word документ {filename}:\n{full_text}'
                    }
                    
                    # Добавляем информацию о хранилище, если файл был сохранен
                    if storage_info:
                        result['storage'] = storage_info
                        result['description'] += f'\n\nФайл сохранен в облаке: {storage_info["url"]}'
                    
                    return result
                else:
                    result = {
                        'type': 'docx',
                        'filename': filename,
                        'description': f'Word документ: {filename} (документ пуст)'
                    }
                    if storage_info:
                        result['storage'] = storage_info
                        result['description'] += f'\n\nФайл сохранен в облаке: {storage_info["url"]}'
                    return result
            except ImportError:
                return {
                    'type': 'docx',
                    'filename': filename,
                    'description': f'Word документ: {filename} (библиотека python-docx не установлена)'
                }
            except Exception as e:
                return {
                    'type': 'docx',
                    'filename': filename,
                    'description': f'Word документ: {filename} (ошибка обработки: {str(e)})'
                }
        
        # Process Excel files (.xlsx)
        elif content_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet' or file_extension in ('xlsx', 'xls'):
            try:
                import openpyxl
                
                # Extract data from Excel
                workbook = openpyxl.load_workbook(io.BytesIO(file_data), data_only=True)
                excel_text = []
                
                for sheet_name in workbook.sheetnames[:3]:  # Limit to first 3 sheets
                    sheet = workbook[sheet_name]
                    excel_text.append(f'\nЛист "{sheet_name}":')
                    
                    for row_idx, row in enumerate(sheet.iter_rows(values_only=True), 1):
                        if row_idx > 50:  # Limit rows per sheet
                            excel_text.append('[... строки пропущены ...]')
                            break
                        row_data = [str(cell) if cell is not None else '' for cell in row]
                        if any(row_data):  # Skip empty rows
                            excel_text.append(' | '.join(row_data))
                
                full_text = '\n'.join(excel_text)
                if len(full_text) > 5000:
                    full_text = full_text[:5000] + '\n\n[... данные обрезаны ...]'
                
                result = {
                    'type': 'excel',
                    'filename': filename,
                    'content': full_text,
                    'description': f'Excel файл {filename}:\n{full_text}'
                }
                
                # Добавляем информацию о хранилище, если файл был сохранен
                if storage_info:
                    result['storage'] = storage_info
                    result['description'] += f'\n\nФайл сохранен в облаке: {storage_info["url"]}'
                
                return result
            except ImportError:
                return {
                    'type': 'excel',
                    'filename': filename,
                    'description': f'Excel файл: {filename} (библиотека openpyxl не установлена)'
                }
            except Exception as e:
                return {
                    'type': 'excel',
                    'filename': filename,
                    'description': f'Excel файл: {filename} (ошибка обработки: {str(e)})'
                }
        
        # Process text files
        elif content_type.startswith('text/') or file_extension in ('txt', 'md', 'py', 'js', 'html', 'css', 'json', 'xml', 'csv', 'log', 'yaml', 'yml'):
            try:
                text_content = file_data.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    text_content = file_data.decode('latin-1')
                except:
                    text_content = f'[Не удалось прочитать содержимое файла {filename}]'
            
            # Limit text length
            if len(text_content) > 5000:
                text_content = text_content[:5000] + '\n\n[... текст обрезан ...]'
            
            result = {
                'type': 'text',
                'filename': filename,
                'content': text_content,
                'description': f'Текстовый файл {filename}:\n{text_content}'
            }
            
            # Добавляем информацию о хранилище, если файл был сохранен
            if storage_info:
                result['storage'] = storage_info
                result['description'] += f'\n\nФайл сохранен в облаке: {storage_info["url"]}'
            
            return result
        
        # Other file types
        else:
            result = {
                'type': 'other',
                'filename': filename,
                'content_type': content_type,
                'description': f'Файл: {filename} (тип: {content_type}) - содержимое не может быть обработано автоматически'
            }
            if storage_info:
                result['storage'] = storage_info
                result['description'] += f'\n\nФайл сохранен в облаке: {storage_info["url"]}'
            return result
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {
            'type': 'error',
            'filename': file.filename if hasattr(file, 'filename') else 'unknown',
            'error': str(e),
            'description': f'Ошибка обработки файла {file.filename if hasattr(file, "filename") else "unknown"}: {str(e)}'
        }


@app.route('/api/chat/with-files', methods=['POST'])
@app.route('/chat/with-files', methods=['POST'])
def chat_with_files():
    """Chat endpoint with file upload support"""
    try:
        # Get form data
        message = request.form.get('message', '')
        provider_name = request.form.get('provider', 'openai')
        
        try:
            temperature = float(request.form.get('temperature', 0.7))
        except (ValueError, TypeError):
            temperature = 0.7
        
        try:
            max_tokens = int(request.form.get('maxTokens', 2000))
        except (ValueError, TypeError):
            max_tokens = 2000
        
        # Get user info
        user_id = request.form.get('user_id')
        user_name = request.form.get('user_first_name')
        username = request.form.get('user_username')
        
        # Process uploaded files
        file_descriptions = []
        file_keys = [key for key in request.files.keys() if key.startswith('file_')]
        file_keys.sort()  # Sort to maintain order
        
        for file_key in file_keys:
            file = request.files[file_key]
            if file and file.filename:
                file_info = process_file(file)
                file_descriptions.append(file_info['description'])
        
        # Combine message with file descriptions
        if file_descriptions:
            files_text = '\n\nПрикрепленные файлы:\n' + '\n'.join([f'- {desc}' for desc in file_descriptions])
            full_message = message + files_text if message else files_text
        else:
            full_message = message
        
        if not full_message.strip():
            return jsonify({
                'success': False,
                'error': 'Message or files are required'
            }), 400
        
        # Get or create chat session
        chat_id = None
        messages_history = []
        if user_id:
            chat_id = get_or_create_chat(user_id, user_name, username, provider_name)
            # Load chat history
            history = get_chat_history(chat_id)
            # Convert to format expected by providers
            messages_history = [
                {"role": msg['role'], "content": msg['content']}
                for msg in history
            ]
        
        provider = get_provider(provider_name)
        
        # Add current user message to history
        messages_history.append({"role": "user", "content": full_message})
        
        # Save user message to database
        if chat_id and user_id:
            try:
                save_message(chat_id, "user", full_message, provider_name, temperature, max_tokens)
            except Exception as db_error:
                print(f"WARNING: Failed to save user message to database: {str(db_error)}")
        
        def generate():
            full_response = ""
            chunk_count = 0
            try:
                print(f"DEBUG: Starting stream with files for provider {provider_name}, message length: {len(full_message)}")
                for chunk in provider.stream(
                    full_message,
                    temperature=temperature,
                    max_tokens=max_tokens,
                    messages=messages_history if messages_history else None
                ):
                    if chunk:
                        chunk_count += 1
                        full_response += chunk
                        try:
                            yield f"data: {json.dumps({'content': chunk})}\n\n"
                        except Exception as yield_error:
                            print(f"ERROR: Failed to yield chunk {chunk_count}: {str(yield_error)}")
                            # Продолжаем обработку даже если не удалось отправить один чанк
                            continue
                
                print(f"DEBUG: Stream with files completed. Total chunks: {chunk_count}, total length: {len(full_response)}")
                
                # Save assistant response to database
                if chat_id and user_id and full_response:
                    try:
                        save_message(chat_id, "assistant", full_response, provider_name, temperature, max_tokens)
                    except Exception as db_error:
                        print(f"WARNING: Failed to save message to database: {str(db_error)}")
                
                yield "data: [DONE]\n\n"
            except ValueError as e:
                # Handle ValueError (API errors from providers)
                error_msg = str(e)
                error_type = type(e).__name__
                print(f"ERROR: Stream generation failed - {error_type}: {error_msg}")
                # Send user-friendly error message (already formatted by provider)
                yield f"data: {json.dumps({'error': error_msg, 'type': error_type})}\n\n"
                yield "data: [DONE]\n\n"
            except GeneratorExit:
                # Клиент закрыл соединение, это нормально
                print(f"DEBUG: Client closed connection. Processed {chunk_count} chunks, total length: {len(full_response)}")
                raise
            except Exception as e:
                error_msg = str(e)
                error_type = type(e).__name__
                print(f"ERROR: Stream generation failed - {error_type}: {error_msg}")
                print(f"DEBUG: Processed {chunk_count} chunks before error, total length: {len(full_response)}")
                import traceback
                traceback.print_exc()
                
                # Если уже была отправлена часть ответа, сохраняем её
                if chat_id and user_id and full_response:
                    try:
                        save_message(chat_id, "assistant", full_response, provider_name, temperature, max_tokens)
                    except Exception as db_error:
                        print(f"WARNING: Failed to save partial message to database: {str(db_error)}")
                
                # Provide user-friendly message for quota/balance errors
                if 'insufficient balance' in error_msg.lower() or 'insufficient_balance' in error_msg.lower() or '402' in error_msg:
                    user_msg = "DeepSeek API: Недостаточно средств на балансе. Пожалуйста, пополните ваш аккаунт на https://platform.deepseek.com"
                elif '413' in error_msg or 'Request too large' in error_msg.lower() or 'tokens per minute' in error_msg.lower():
                    # Ошибка 413 от Groq - запрос слишком большой
                    user_msg = "Запрос слишком большой. Система автоматически разбивает его на части. Если ошибка повторяется, попробуйте сократить размер сообщения или истории."
                elif 'quota' in error_msg.lower() or 'insufficient_quota' in error_msg.lower() or error_type == 'RateLimitError':
                    user_msg = "Превышена квота использования API. Пожалуйста, проверьте ваш план и настройки биллинга."
                else:
                    user_msg = f"Ошибка при генерации ответа: {error_msg}"
                
                try:
                    yield f"data: {json.dumps({'error': user_msg, 'type': error_type, 'details': error_msg})}\n\n"
                    yield "data: [DONE]\n\n"
                except Exception as yield_error:
                    print(f"ERROR: Failed to yield error message: {str(yield_error)}")
                    # Если не удалось отправить ошибку, просто завершаем
                    pass
        
        return Response(
            stream_with_context(generate()),
            mimetype='text/event-stream',
            headers={
                'Cache-Control': 'no-cache',
                'X-Accel-Buffering': 'no'
            }
        )
    
    except ValueError as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 400
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/chat/history', methods=['GET'])
@app.route('/chat/history', methods=['GET'])
def get_chat_history_endpoint():
    """Get chat history for user"""
    try:
        user_id = request.args.get('user_id')
        chat_id = request.args.get('chat_id')
        limit = int(request.args.get('limit', 50))
        
        if not user_id and not chat_id:
            return jsonify({
                'success': False,
                'error': 'user_id or chat_id is required'
            }), 400
        
        if chat_id:
            # Get specific chat with messages
            chat = get_chat_with_messages(int(chat_id))
            if not chat:
                return jsonify({
                    'success': False,
                    'error': 'Chat not found'
                }), 404
            
            return jsonify({
                'success': True,
                'data': chat
            })
        else:
            # Get user's chats
            chats = get_user_chats(user_id, limit=10)
            return jsonify({
                'success': True,
                'data': {
                    'chats': chats
                }
            })
    
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/chat/<int:chat_id>/messages', methods=['GET'])
@app.route('/chat/<int:chat_id>/messages', methods=['GET'])
def get_chat_messages(chat_id):
    """Get messages for a specific chat"""
    try:
        limit = int(request.args.get('limit', 50))
        messages = get_chat_history(chat_id, limit)
        
        return jsonify({
            'success': True,
            'data': {
                'chat_id': chat_id,
                'messages': messages
            }
        })
    
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


if __name__ == '__main__':
    # Development server only - use gunicorn for production
    port = int(os.environ.get('PORT', 8000))
    debug_mode = os.environ.get('FLASK_DEBUG', 'False').lower() == 'true'
    app.run(host='0.0.0.0', port=port, debug=debug_mode)